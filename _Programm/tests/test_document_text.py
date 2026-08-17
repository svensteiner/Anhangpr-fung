# -*- coding: utf-8 -*-
"""Tests für den format-agnostischen Seiten-Loader (Word-Konnektor)."""

import docx
import pytest

from anhangspruefer.parsers.document_text import load_page_texts
from anhangspruefer.vorjahresvergleich.extractor import extract_items
from anhangspruefer.vorjahresvergleich.comparator import opening_value, closing_value


def _stand_spiegel_docx(path):
    """Kleiner Anhang-Auszug: Text + Stand-Spiegel als Word-Tabelle."""
    doc = docx.Document()
    doc.add_paragraph("Zusammensetzung und Entwicklung der Rückstellungen:")
    t = doc.add_table(rows=3, cols=5)
    hdr = t.rows[0].cells
    hdr[1].text = "Stand 01.01.2025"
    hdr[2].text = "Verwendung"
    hdr[3].text = "Zuweisung"
    hdr[4].text = "Stand 31.12.2025"
    for c in t.rows[1].cells[1:]:   # reine Währungszeile
        c.text = "€"
    row = t.rows[2].cells
    row[0].text = "Rückstellungen für Abfertigungen"
    row[1].text = "103.618,02"
    row[2].text = "0,00"
    row[3].text = "76.456,98"
    row[4].text = "180.075,00"
    doc.save(str(path))


def test_load_docx_paragraphs_and_table(tmp_path):
    p = tmp_path / "anhang.docx"
    _stand_spiegel_docx(p)

    pages = load_page_texts(p)
    assert isinstance(pages, list) and pages

    joined = "\n".join(pages)
    # Absatz erhalten
    assert any("Zusammensetzung" in pg for pg in pages)
    # Tabellenkopf + Werte linearisiert vorhanden
    assert "Stand 01.01.2025" in joined
    assert "103.618,02" in joined
    assert "180.075,00" in joined
    # reine Währungszeile ("€ …") wurde verworfen
    assert "€" not in joined
    # Tabelle als eigene "Seite" isoliert (nicht mit dem Absatz vermischt)
    table_pages = [pg for pg in pages if "Stand 01.01.2025" in pg]
    assert len(table_pages) == 1
    assert "Zusammensetzung" not in table_pages[0]


def test_extract_items_from_docx_stand_spiegel(tmp_path):
    """Der Standard-Extraktor liest den Word-Stand-Spiegel korrekt als
    Eröffnungs-/Schlusswert (Bilanzkontinuität)."""
    p = tmp_path / "anhang.docx"
    _stand_spiegel_docx(p)

    items = extract_items(p)
    hit = [it for it in items if "Abfertigungen" in it.label]
    assert hit, f"Posten nicht gefunden in {[it.label for it in items]}"
    it = hit[0]
    assert opening_value(it) == pytest.approx(103618.02)   # Stand 01.01.
    assert closing_value(it) == pytest.approx(180075.00)   # Stand 31.12.


def test_pdf_suffix_dispatch_is_case_insensitive(tmp_path, monkeypatch):
    """`.PDF` (Großschreibung) landet beim PDF-Reader."""
    called = {}

    def fake_pdf(path, x_tolerance=2):
        called["path"] = path
        return ["seite"]

    monkeypatch.setattr(
        "anhangspruefer.parsers.document_text._pdf_page_texts", fake_pdf
    )
    out = load_page_texts(tmp_path / "Bericht.PDF")
    assert out == ["seite"]
    assert called  # PDF-Reader wurde aufgerufen


def test_unsupported_suffix_raises(tmp_path):
    p = tmp_path / "x.txt"
    p.write_text("hallo", encoding="utf-8")
    with pytest.raises(ValueError):
        load_page_texts(p)


def test_docx_exclude_tables_for_text_compare(tmp_path):
    """include_tables=False liefert nur Fließtext (für den Textvergleich),
    damit unterschiedlich linearisierte Word/PDF-Tabellen keine
    Scheinänderungen erzeugen."""
    p = tmp_path / "anhang.docx"
    _stand_spiegel_docx(p)

    with_tables = load_page_texts(p, include_tables=True)
    without_tables = load_page_texts(p, include_tables=False)

    assert any("103.618,02" in pg for pg in with_tables)      # Zahl in Tabelle
    assert all("103.618,02" not in pg for pg in without_tables)  # Tabelle raus
    assert any("Zusammensetzung" in pg for pg in without_tables)  # Absatz bleibt


# ---------------------------------------------------------------------------
# Textvergleich-Heuristik (Tabellen fern vom Fließtext halten)
# ---------------------------------------------------------------------------
from anhangspruefer.vorjahresvergleich.text_compare import (
    _is_table_header_line,
    _is_narrative,
)


def test_table_header_lines_detected():
    assert _is_table_header_line("Stand Stand")
    assert _is_table_header_line("Aktiv Passiv Aktiv Passiv Bewegung")
    assert _is_table_header_line("Nutzungsdauer in Jahren")
    assert _is_table_header_line("31.12.2024 31.12.2023")
    assert not _is_table_header_line(
        "Der Jahresabschluss wurde nach den Vorschriften der §§ 189 ff UGB aufgestellt."
    )


def test_narrative_rejects_table_rows():
    # Währungszeichen bzw. >= 3 Beträge => Tabellenzeile, kein Fließtext
    assert not _is_narrative(
        "Sonstige Verbindlichkeiten € Entwicklung eines webbasierten Workflows 2.800,00"
    )
    assert not _is_narrative(
        "Unternehmensberatung Gmbh München -96.364,29 85,00 -17.692,21 Musterfirma"
    )
    # echter Fließtext bleibt Fließtext (auch mit ein/zwei Beträgen)
    assert _is_narrative(
        "Die Forderungen und sonstigen Vermögensgegenstände wurden mit dem Nennwert angesetzt."
    )
    assert _is_narrative(
        "Die Aufwendungen für den Abschlussprüfer betragen EUR 8.000,00 (Vorjahr: 7.250,00)."
    )

# -*- coding: utf-8 -*-
"""Tests für die Syngroup-Pipeline und den wertbasierten Duplikat-Match."""

from anhangspruefer.pipelines import get_pipeline
from anhangspruefer.pipelines.syngroup import SyngroupPipeline, _is_continuity_item
from anhangspruefer.vorjahresvergleich.extractor import AnhangItem
from anhangspruefer.vorjahresvergleich import comparator as C


def _item(label, current_values, prior_values=None):
    return AnhangItem(
        label=label,
        page=1,
        current_values=current_values,
        prior_values=prior_values or [],
    )


# ---------------------------------------------------------------------------
# Mandanten-Auswahl
# ---------------------------------------------------------------------------
def test_get_pipeline_selects_syngroup():
    assert get_pipeline("Syngroup Management Consulting AG").name == "syngroup"
    assert get_pipeline("syngroup").name == "syngroup"
    # unbekannt -> Standard
    assert get_pipeline("Irgendwer GmbH").name == "standard"


# ---------------------------------------------------------------------------
# Filter der vergleichbaren Kontinuitätsposten
# ---------------------------------------------------------------------------
def test_syngroup_filter_drops_latente_steuern():
    # Latente-Steuern-Matrixzeilen (keine Bilanzkontinuität)
    assert not _is_continuity_item(_item("Aktivposten", [8014.44, 2925.38, 5089.06]))
    assert not _is_continuity_item(_item("Rückstellung für Abfertigung", [16703.24, 16703.24, 0.0]))
    assert not _is_continuity_item(_item("Aktive (+)/ passive (-) latente Steuerabgrenzung", [5685.07]))


def test_syngroup_filter_drops_document_noise():
    assert not _is_continuity_item(_item("Software Nutzungsdauer in Jahren 3,00 -", [4.0]))
    assert not _is_continuity_item(_item("1010 Wien, Kärntner Ring 17/", [17.0]))
    assert not _is_continuity_item(_item("119149 Seite", [8.0]))
    assert not _is_continuity_item(_item("Datum, Unterschriften der Vorstände", [10.0]))


def test_syngroup_filter_drops_year_artifacts_and_breakdown():
    # "RST JAB 2024 10.000,00" -> current=[2024, 10000]: 2024 ist ein Jahr-Artefakt
    assert not _is_continuity_item(_item("RST JAB", [2024.0, 10000.0]))
    # Zwei-Spalten-Split: die Jahreszahl landet in current, der Betrag in prior
    assert not _is_continuity_item(_item("Kooperationsbetrag", [2024.0], [7300.0]))
    assert not _is_continuity_item(_item("Telefongebühren 12/", [2024.0], [1485.20]))
    # Nutzungsdauer-Rest ("… 9,50 -")
    assert not _is_continuity_item(_item("Bauten auf fremdem Grund 9,50 -", [10.0]))
    # umgebrochene "SUMME RÜCKSTELLUNGEN"
    assert not _is_continuity_item(_item("RÜCKSTELLUNGEN", [755301.13]))


def test_syngroup_filter_keeps_real_positions():
    # Plural = echter Rückstellungsspiegel-Posten (nicht die Steuerlatenz-Zeile)
    assert _is_continuity_item(_item("Rückstellungen für Abfertigungen", [103618.02, 0.0, 180075.0]))
    assert _is_continuity_item(_item("Rückstellung Prämien MA", [261576.86, 0.0, 261576.86]))
    assert _is_continuity_item(_item("Grundkapital", [70000.0]))
    # echte Vorjahres-Only-Posten (kein Jahr-Artefakt) bleiben erhalten
    assert _is_continuity_item(_item("Entwicklung eines webbasierten Workflows", [2800.0]))
    assert _is_continuity_item(_item("Verpflichtungen aus Leasingverträgen", [132197.16, 392630.04]))


def test_pipeline_extract_applies_filter(tmp_path):
    import docx

    doc = docx.Document()
    doc.add_paragraph("Zusammensetzung und Entwicklung der Rückstellungen:")
    t = doc.add_table(rows=2, cols=5)
    hdr = t.rows[0].cells
    hdr[1].text = "Stand 01.01.2025"
    hdr[4].text = "Stand 31.12.2025"
    row = t.rows[1].cells
    row[0].text = "Aktivposten"          # Latente-Steuern-Zeile -> muss gefiltert werden
    row[1].text = "8.014,44"
    row[2].text = "2.925,38"
    row[4].text = "5.089,06"
    p = tmp_path / "anhang.docx"
    doc.save(str(p))

    labels = [it.label for it in SyngroupPipeline().extract_anhang_items(p)]
    assert "Aktivposten" not in labels


# ---------------------------------------------------------------------------
# Wertbasierter Duplikat-Match im Comparator
# ---------------------------------------------------------------------------
def test_find_match_prefers_value_matching_candidate():
    """Zwei gleichnamige Vorjahresposten: der wertgleiche wird gewählt."""
    # z.B. "Investitionsprämie 7%" steht unter Immateriell (0,12->0,00) UND
    # unter Sachanlagen (448,66->368,41); gesucht wird der Eröffnungswert 368,41.
    immateriell = _item("Investitionsprämie 7%", [0.12, 0.12, 0.00])
    sachanlagen = _item("Investitionsprämie 7%", [448.66, 80.25, 368.41])
    index = C._index_by_normalized([immateriell, sachanlagen])
    key = immateriell.label_key_compact

    match, score = C._find_match(key, index, target_value=368.41)
    assert C.closing_value(match) == 368.41
    assert score == 1.0

    # Ohne Zielwert bleibt es beim ersten Treffer (Rückwärtskompatibilität)
    match0, _ = C._find_match(key, index)
    assert match0 is immateriell

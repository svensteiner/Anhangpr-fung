"""
Format-agnostischer Seiten-Text-Loader — der "Konnektor" zwischen den
Roh-Unterlagen eines Mandanten und dem gemeinsamen "Hirn".

Leitidee
========
Die gesamte Vergleichs-/Prüflogik (Zahlen- UND Textvergleich) arbeitet nur mit
EINER Datenform: einer Liste linearisierter Seitentexte (``list[str]``). WIE
diese Seitentexte aus der konkreten Datei gewonnen werden, kapselt allein dieser
Loader:

    * ``.pdf``  -> pdfplumber  (unverändertes bisheriges Verhalten)
    * ``.docx`` -> python-docx  (Absätze + Tabellen in Dokumentreihenfolge)

Dadurch funktioniert die bestehende, für PDF gebaute Heuristik unverändert auch
für Word-Anhänge, ohne dass Extraktor oder Textvergleich das Dateiformat kennen
müssen.

Word-Besonderheit — Tabellen
============================
Der Zahlen-Extraktor erwartet je Zeile "Label  Zahl  Zahl …" (durch Whitespace
getrennte nachlaufende Zahlen). Word-Tabellen werden daher Zeile für Zeile zu
genau dieser Form linearisiert (Zellen durch mehrere Leerzeichen getrennt).

Jede Tabelle wird als EIGENE "Seite" ausgegeben. Das isoliert den seiten-lokalen
Parser-Zustand des Extraktors (Spaltenkopf-/Anlagenspiegel-Erkennung), sodass
sich Zustände nicht von einer Tabelle in die nächste "durchschleppen".

Vollständig lokal. Kein Netzwerk, kein externer Aufruf.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Optional


# Trennschärfe für pdfplumber (siehe extractor.X_TOLERANCE / text_compare).
# 2 rekonstruiert Wortgrenzen über Zeichenabstände zuverlässig.
DEFAULT_X_TOLERANCE = 2

#: Formate, die dieser Konnektor lesen kann. Single Source of Truth – Aufrufer,
#: die eine Datei nur bei Lesbarkeit heranziehen wollen, fragen ``kann_lesen``
#: statt eine eigene Liste zu pflegen.
LESBARE_SUFFIXE = (".pdf", ".docx")

#: Bis zu dieser Zeichenzahl je Seite trägt eine Unterlage praktisch keinen
#: lesbaren Text – typischer Fall: ein Scan ohne Texterkennung (OCR), der 0
#: Zeichen liefert. Bewusst niedrig, damit ein Dokument aus Deckblatt und
#: Trennseiten nicht schon deshalb als leer gilt.
LEER_GRENZE_JE_SEITE = 25

#: Darunter ist die Unterlage auffällig textarm: Teil-Scan, oder ein PDF mit
#: fehlerhaft eingebetteten Schriften. Ein normal gesetzter Jahresabschluss
#: liegt bei 1.500–3.000 Zeichen je Seite.
TEXTARM_GRENZE_JE_SEITE = 250

# Zell-Trenner beim Linearisieren einer Word-Tabellenzeile. Bewusst breit, damit
# der Extraktor Label und Werte sicher voneinander trennen kann.
_CELL_SEP = "   "


def load_page_texts(
    path,
    x_tolerance: int = DEFAULT_X_TOLERANCE,
    include_tables: bool = True,
) -> list[str]:
    """Liefert die Seitentexte einer Unterlage als ``list[str]``.

    Args:
        path: Pfad zur Datei (.pdf oder .docx; Groß-/Kleinschreibung egal).
        x_tolerance: nur für PDF relevant (pdfplumber-Wortabstand).
        include_tables: nur für Word relevant. True (Standard) = Tabellen als
            eigene Seiten mitliefern (für den Zahlenvergleich). False = NUR
            Fließtext-Absätze (für den Textvergleich, damit unterschiedlich
            linearisierte Tabellen aus Word/PDF nicht als "neue/fehlende
            Textteile" auftauchen).

    Returns:
        Liste linearisierter Seitentexte. Bei Word: Absatz-Blöcke und (sofern
        ``include_tables``) Tabellen in Dokumentreihenfolge, jede Tabelle als
        eigener Eintrag.

    Raises:
        ValueError: bei nicht unterstütztem Dateiformat.
    """
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return _pdf_page_texts(path, x_tolerance)
    if suffix == ".docx":
        return _docx_page_texts(path, include_tables=include_tables)
    raise ValueError(
        f"Nicht unterstütztes Dateiformat für die Textextraktion: '{suffix}'. "
        "Unterstützt werden .pdf und .docx."
    )


def kann_lesen(path) -> bool:
    """Kann dieser Konnektor die Datei lesen? (Entscheidung nur am Suffix.)

    Für Aufrufer, die eine Unterlage nur dann als Textquelle heranziehen, wenn
    sie überhaupt lesbar ist – z.B. der interne Abgleich, dem ein Excel-Beleg
    als Bilanz-/GuV-Quelle nichts nützt.
    """
    return Path(path).suffix.lower() in LESBARE_SUFFIXE


@dataclass(frozen=True)
class Textausbeute:
    """Wieviel lesbarer Text steckt in einer Unterlage?

    Hintergrund: Ein gescannter Abschluss ohne Texterkennung liefert leere
    Seitentexte. Die gesamte Prüflogik läuft dann fehlerfrei durch und erzeugt
    ein vollständig formatiertes Arbeitspapier OHNE inhaltliche Grundlage –
    fachlich die gefährlichste Fehlnutzung, weil sie wie ein Ergebnis aussieht.
    """

    datei: str
    seiten: int
    zeichen: int

    @property
    def zeichen_je_seite(self) -> float:
        return self.zeichen / self.seiten if self.seiten else 0.0

    @property
    def ist_leer(self) -> bool:
        """Kein verwertbarer Text – die Prüfung darf nicht durchlaufen."""
        return self.seiten == 0 or self.zeichen_je_seite <= LEER_GRENZE_JE_SEITE

    @property
    def ist_textarm(self) -> bool:
        """Verwertbar, aber verdächtig wenig – Warnung, kein Abbruch."""
        return not self.ist_leer and self.zeichen_je_seite < TEXTARM_GRENZE_JE_SEITE

    @property
    def meldung(self) -> str:
        """Ein Satz für den Anwender; leer, wenn die Unterlage unauffällig ist."""
        if self.seiten == 0:
            return (f"„{self.datei}“ enthält keine lesbaren Seiten. Bitte prüfen, "
                    f"ob die Datei vollständig und unbeschädigt ist.")
        if self.ist_leer:
            return (f"„{self.datei}“ enthält keinen lesbaren Text "
                    f"({self.seiten} Seiten, {self.zeichen} Zeichen). Das ist "
                    f"vermutlich ein Scan ohne Texterkennung. Bitte das PDF aus "
                    f"dem Erstellungsprogramm mit Textebene anfordern oder eine "
                    f"Texterkennung (OCR) darüber laufen lassen.")
        if self.ist_textarm:
            return (f"„{self.datei}“ enthält auffällig wenig Text (rund "
                    f"{self.zeichen_je_seite:.0f} Zeichen je Seite auf "
                    f"{self.seiten} Seiten). Möglicherweise ist ein Teil des "
                    f"Dokuments nur gescannt. Das Ergebnis kann unvollständig "
                    f"sein – bitte stichprobenweise gegen das Dokument prüfen.")
        return ""


def textausbeute(page_texts: list[str], datei: str) -> Textausbeute:
    """Misst die Textausbeute bereits gelesener Seitentexte."""
    return Textausbeute(
        datei=datei,
        seiten=len(page_texts),
        zeichen=sum(len(t or "") for t in page_texts),
    )


def pruefe_textausbeute(path, anzeigename: Optional[str] = None) -> Optional[Textausbeute]:
    """Liest eine Unterlage und misst, wieviel Text sie trägt.

    Args:
        path: Pfad zur Datei.
        anzeigename: Name für die Anwendermeldung (die hochgeladene Datei liegt
            unter einem technischen Temp-Namen). Ohne Angabe der Dateiname.

    Returns:
        ``Textausbeute`` – oder ``None``, wenn die Datei mit diesem Konnektor
        gar nicht gelesen werden kann (fremdes Format, beschädigt). Dann
        entscheidet der aufrufende Modus wie bisher weiter: die Prüfung soll
        keinen Weg blockieren, der heute funktioniert.
    """
    try:
        pages = load_page_texts(path)
    except Exception:
        return None
    return textausbeute(pages, anzeigename or Path(path).name)


def _pdf_page_texts(path, x_tolerance: int = DEFAULT_X_TOLERANCE) -> list[str]:
    """Seitentexte aus einem PDF (pdfplumber) — unverändertes bisheriges Verhalten."""
    import pdfplumber

    with pdfplumber.open(str(path)) as pdf:
        return [p.extract_text(x_tolerance=x_tolerance) or "" for p in pdf.pages]


def _linearize_row(cells: list[str]) -> str:
    """Wandelt eine Word-Tabellenzeile in "Label   Zahl   Zahl …" um.

    Leere Zeilen sowie reine Trenn-/Währungszeilen (z.B. "€ € € €") liefern "",
    damit sie den Extraktor nicht als Pseudo-Label stören.
    """
    joined = _CELL_SEP.join(c for c in cells).strip()
    # Ohne Buchstaben/Ziffern trägt die Zeile keine Information (nur €, Striche…).
    if not any(ch.isalnum() for ch in joined):
        return ""
    return joined


def _docx_page_texts(path, include_tables: bool = True) -> list[str]:
    """Seitentexte aus einem Word-Dokument (python-docx).

    Absätze und Tabellen werden in ihrer echten Dokumentreihenfolge gelesen
    (dafür wird der Body direkt iteriert — ``document.paragraphs`` und
    ``document.tables`` allein verlieren die Verschachtelung). Zusammenhängende
    Absätze bilden je einen Seiten-Block; jede Tabelle wird als eigene Seite
    ausgegeben. Bei ``include_tables=False`` werden Tabellen übersprungen
    (nur Fließtext – für den Textvergleich).
    """
    import docx
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(str(path))
    pages: list[str] = []
    para_buf: list[str] = []

    def flush_paragraphs() -> None:
        if para_buf:
            pages.append("\n".join(para_buf))
            para_buf.clear()

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            text = Paragraph(child, doc).text.strip()
            if text:
                para_buf.append(text)
        elif isinstance(child, CT_Tbl):
            # Tabelle beendet den laufenden Absatz-Block.
            flush_paragraphs()
            if not include_tables:
                continue
            table = Table(child, doc)
            lines: list[str] = []
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                line = _linearize_row(cells)
                if line:
                    lines.append(line)
            if lines:
                pages.append("\n".join(lines))

    flush_paragraphs()
    return pages
